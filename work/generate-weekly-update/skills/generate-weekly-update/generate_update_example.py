"""
Example script demonstrating the weekly update document generation.
Uses placeholder data — replace with actual content when running.
"""

from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def add_bullet(paragraph, level=0, num_id=1):
    pPr = paragraph._element.get_or_add_pPr()
    numPr = parse_xml(r'<w:numPr %s><w:ilvl w:val="%d"/><w:numId w:val="%d"/></w:numPr>' % (nsdecls('w'), level, num_id))
    pPr.insert(0, numPr)


def set_font(paragraph, font_name='Arial', font_size=11, bold=False):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold


def add_paragraph_with_font(doc, text, bold=False):
    p = doc.add_paragraph(text)
    set_font(p, 'Arial', 11, bold)
    return p


def add_bullet_paragraph(doc, text, level=0, num_id=1):
    p = doc.add_paragraph(text)
    add_bullet(p, level, num_id)
    set_font(p, 'Arial', 11)
    return p


# Load last week's document as template
doc = Document("previous-week-update.docx")

# Clear all paragraphs
for paragraph in doc.paragraphs[:]:
    p = paragraph._element
    p.getparent().remove(p)

# Title
add_paragraph_with_font(doc, "Update for 1st March - 5th March week")
doc.add_paragraph()

# Summary
add_paragraph_with_font(doc, "Summary", bold=True)
add_paragraph_with_font(doc, "[TO BE GENERATED AFTER REVIEW]")
doc.add_paragraph()
doc.add_paragraph()

# System / Process / Tools
add_paragraph_with_font(doc, "System / Process / Tools", bold=True)
doc.add_paragraph()
add_bullet_paragraph(doc, "This week, the primary focus was on [describe main focus area for the week].", level=0, num_id=1)
doc.add_paragraph()
doc.add_paragraph()

# Good
add_paragraph_with_font(doc, "Good", bold=True)
doc.add_paragraph()

add_bullet_paragraph(doc, "Client A Performance - [Owner Names]", level=0, num_id=2)
add_bullet_paragraph(doc, "Collections performance trending positive at $XXK", level=1, num_id=2)
add_bullet_paragraph(doc, "Positive increase in contact rate by X%", level=1, num_id=2)

add_bullet_paragraph(doc, "Client B Launch - [Owner Names]", level=0, num_id=2)
add_bullet_paragraph(doc, "Successfully launched on [date]", level=1, num_id=2)

add_bullet_paragraph(doc, "Client C Allocation Increase - [Owner Names]", level=0, num_id=2)
add_bullet_paragraph(doc, "Positive increment in collections resulting in favorable client review", level=1, num_id=2)
add_bullet_paragraph(doc, "Client agreed to increase allocation to XXK from XXK", level=1, num_id=2)

doc.add_paragraph()

# Bad
add_paragraph_with_font(doc, "Bad", bold=True)
doc.add_paragraph()

add_bullet_paragraph(doc, "Client D Performance Below Target - [Owner Names]", level=0, num_id=3)
add_bullet_paragraph(doc, "Payment collection has not met expectations on the new account set", level=1, num_id=3)

add_bullet_paragraph(doc, "Vendor Integration Issues - [Owner Names]", level=0, num_id=3)
add_bullet_paragraph(doc, "Vendor account suspension impacting campaigns", level=1, num_id=3)
add_bullet_paragraph(doc, "Delay in vendor onboarding and integrations for Client E", level=1, num_id=3)

doc.add_paragraph()

# Ugly
add_paragraph_with_font(doc, "Ugly", bold=True)
doc.add_paragraph()

add_bullet_paragraph(doc, "Compliance Issue on Client A - [Owner Names]", level=0, num_id=4)
add_bullet_paragraph(doc, "Approximately X incorrect communications sent with wrong details", level=1, num_id=4)

add_bullet_paragraph(doc, "Low Collections on Recently Launched Accounts - [Owner Names]", level=0, num_id=4)
add_bullet_paragraph(doc, "Client F - Collections at $XK, expected $XK", level=1, num_id=4)
add_bullet_paragraph(doc, "Client G - No payments collected this week", level=1, num_id=4)

add_bullet_paragraph(doc, "Reporting Failure - [Owner Names]", level=0, num_id=4)
add_bullet_paragraph(doc, "Incident resulting in incorrect end-of-day report sent to client, requiring escalation and resend", level=1, num_id=4)

doc.add_paragraph()

# Support Needed
add_paragraph_with_font(doc, "Support Needed", bold=True)
doc.add_paragraph()

add_bullet_paragraph(doc, "Bot Improvements", level=0, num_id=5)
add_bullet_paragraph(doc, "Improvement in reporting efficiency", level=1, num_id=5)
add_bullet_paragraph(doc, "Accurate disposition mapping", level=1, num_id=5)

add_bullet_paragraph(doc, "2-way SMS integration for [platform]", level=0, num_id=5)

doc.add_paragraph()

# Pre-Sales
add_paragraph_with_font(doc, "Pre-Sales / Sales-handover Stage", bold=True)
doc.add_paragraph()

add_bullet_paragraph(doc, "Conducted pre-sales discussion for Prospect A", level=0, num_id=6)
add_bullet_paragraph(doc, "Value solutioning planned with Prospect B on [date]", level=0, num_id=6)
add_bullet_paragraph(doc, "Created prototype for new product, defined discovery questions", level=0, num_id=6)

doc.add_paragraph()

# Partners
add_paragraph_with_font(doc, "Partners / Vendors", bold=True)
doc.add_paragraph()

add_bullet_paragraph(doc, "Vendor Connection Challenges for Client E - [Owner Names]", level=0, num_id=7)
add_bullet_paragraph(doc, "Unable to complete vendor onboarding despite multiple attempts", level=1, num_id=7)

doc.add_paragraph()
doc.add_paragraph()

# POPW
add_paragraph_with_font(doc, "Progress on the Previous Week's Bad / Ugly", bold=True)
doc.add_paragraph()

add_bullet_paragraph(doc, "Client H Go-Live - [Owner Names]", level=0, num_id=10)
add_bullet_paragraph(doc, "Client delayed call this week, pending items on their end", level=1, num_id=10)

add_bullet_paragraph(doc, "Client I Go-Live - [Owner Names]", level=0, num_id=10)
add_bullet_paragraph(doc, "Timeline for account setup pending confirmation", level=1, num_id=10)
add_bullet_paragraph(doc, "Efforts underway to renegotiate scope for phased go-live", level=1, num_id=10)

# Save
doc.save("Update-1-Mar-5-Mar-DRAFT.docx")
print("Document generated successfully!")
